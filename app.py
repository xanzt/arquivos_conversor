import streamlit as st
from PyPDF2 import PdfReader
from docx import Document
from fpdf import FPDF
from tabula import read_pdf
import pandas as pd
import io


def pdf_to_word(pdf_file):
    pdf_reader = PdfReader(pdf_file)
    doc = Document()

    for page in pdf_reader.pages:
        text = page.extract_text()
        if text:
            doc.add_paragraph(text)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


def excel_to_word(excel_file):
    # Carregar o arquivo Excel diretamente do BytesIO
    df = pd.read_excel(excel_file, engine='openpyxl')
    doc = Document()

    # Criar uma tabela no documento Word com a quantidade de linhas e colunas adequadas
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'

    # Adiciona os títulos das colunas na primeira linha da tabela
    for i, column in enumerate(df.columns):
        table.cell(0, i).text = str(column)

    # Preenche a tabela com os dados das linhas do DataFrame
    for row in df.itertuples(index=False, name=None):
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)

    # Salva o arquivo gerado no formato Word em memória
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


def excel_to_pdf(excel_file):
    df = pd.read_excel(excel_file)
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=10)

    col_widths = [max(len(str(value)) for value in df[column]) * 5 for column in df.columns]

    for i, column in enumerate(df.columns):
        pdf.cell(col_widths[i], 10, str(column), border=1, align='C')
    pdf.ln()

    for row in df.values:
        for i, value in enumerate(row):
            pdf.cell(col_widths[i], 10, str(value), border=1, align='C')
        pdf.ln()

    output = io.BytesIO()
    pdf_output = pdf.output(dest='S').encode('latin1')
    output.write(pdf_output)
    output.seek(0)
    return output.getvalue()


def pdf_to_excel(pdf_file):
    try:
        dfs = read_pdf(pdf_file, pages="all", multiple_tables=True, stream=True)
        output = io.BytesIO()

        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            for i, df in enumerate(dfs):
                df.to_excel(writer, index=False, sheet_name=f"Tabela_{i + 1}")

        output.seek(0)
        return output.getvalue()
    except Exception as e:
        st.error(f"Erro ao processar o PDF: {e}")
        return None


def main():
    logo_url = "https://i.ibb.co/qBwfjTX/LOGO-BRANCO.png"
    st.image(logo_url, width=200)

    st.title("Conversor de Arquivos")

    st.sidebar.header("Opções de Conversão")
    option = st.sidebar.selectbox(
        "Escolha a conversão:",
        ("PDF para Word", "Excel para Word", "Excel para PDF", "PDF para Excel")
    )

    # Adicionando o texto abaixo da caixa de seleção
    st.sidebar.markdown(
        """
        <p style="font-size: 12px; color: gray; margin-top: 350px;">
        By: Alexandre Carvalho.
        </p>
        """, unsafe_allow_html=True
    )

    uploaded_file = st.file_uploader("Envie o arquivo", type=["pdf", "xlsx"])

    if uploaded_file is not None:
        if option == "PDF para Word":
            if uploaded_file.name.endswith(".pdf"):
                st.success("Arquivo PDF carregado!")
                word_data = pdf_to_word(uploaded_file)
                st.download_button(
                    label="Baixar Word",
                    data=word_data,
                    file_name="convertido.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.error("Por favor, envie um arquivo PDF.")

        elif option == "Excel para Word":
            if uploaded_file.name.endswith(".xlsx"):
                st.success("Arquivo Excel carregado!")

                # Certificar-se de que o arquivo foi processado corretamente
                word_data = excel_to_word(uploaded_file)

                if word_data:
                    st.download_button(
                        label="Baixar Word",
                        data=word_data,
                        file_name="convertido.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                else:
                    st.error("Erro ao converter o arquivo Excel para Word.")
            else:
                st.error("Por favor, envie um arquivo Excel (.xlsx).")

        elif option == "Excel para PDF":
            if uploaded_file.name.endswith(".xlsx"):
                st.success("Arquivo Excel carregado!")
                pdf_data = excel_to_pdf(uploaded_file)
                st.download_button(
                    label="Baixar PDF",
                    data=pdf_data,
                    file_name="convertido.pdf",
                    mime="application/pdf"
                )
            else:
                st.error("Por favor, envie um arquivo Excel (.xlsx).")

        elif option == "PDF para Excel":
            if uploaded_file.name.endswith(".pdf"):
                st.success("Arquivo PDF carregado!")
                excel_data = pdf_to_excel(uploaded_file)
                if excel_data:
                    st.download_button(
                        label="Baixar Excel",
                        data=excel_data,
                        file_name="convertido.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )
            else:
                st.error("Por favor, envie um arquivo PDF.")

    # Adicionando a marca d'água "by: Alexandre Carvalho" no canto inferior esquerdo
    st.markdown(
        """
        <style>
        #watermark {
            position: fixed;
            bottom: 10px;
            left: 10px;
            font-size: 12px;
            color: gray;
            opacity: 0.6;
        }
        </style>
        <div id="watermark">by: Alexandre Carvalho</div>
        """, unsafe_allow_html=True
    )


if __name__ == "__main__":
    main()
